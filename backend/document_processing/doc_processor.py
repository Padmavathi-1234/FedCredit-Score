"""
DOCX Processor — extracts paragraphs and tables from Word documents.
"""

import re
from typing import Dict, Any

try:
    from docx import Document
except ImportError:
    Document = None


def process_docx(file_path: str) -> Dict[str, Any]:
    """Parse a DOCX file and extract financial metrics from text and tables."""
    if Document is None:
        return {"error": "python-docx is required for DOCX processing"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": f"Cannot open DOCX file: {e}"}

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # Also pull tables
    table_data: list = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_data.append(rows)
        # Add table text to full_text for metric extraction
        for row in rows:
            full_text += " " + " ".join(row)

    # Reuse the financial-value extractor
    from backend.document_processing.pdf_processor import _extract_financial_values

    extracted = _extract_financial_values(full_text)
    extracted["raw_text"] = full_text[:3000]
    extracted["tables"] = table_data[:5]
    extracted["source"] = "docx"
    return extracted
