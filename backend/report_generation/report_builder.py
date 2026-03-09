"""
Report Builder — generates a branded PDF Credit Appraisal Memorandum (CAM)
using python-docx and docx2pdf.
"""

import os
import io
import re
import tempfile
import uuid
import logging
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

def _add_markdown_para(doc, text: str, bold_all: bool = False, italic: bool = False):
    """
    Add a paragraph to the document, converting markdown `**bold**` into bold runs.
    """
    p = doc.add_paragraph()
    if not text:
        return p
        
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        
        # Odd indices are the captured bold groups
        if i % 2 == 1 or bold_all:
            run.bold = True
            
        if italic:
            run.italic = True
    return p

def _create_styled_table(doc, rows_data, col_widths=None):
    """Helper to create a shaded, grid-styled table."""
    if not rows_data:
        return None
        
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[r_idx].cells
        for c_idx, cell_data in enumerate(row_data):
            # Write text safely
            cell_text = str(cell_data) if cell_data is not None else "N/A"
            row_cells[c_idx].text = cell_text
            
            # Make header row bold
            if r_idx == 0:
                for run in row_cells[c_idx].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
    return table

def generate_report(analysis_result: Dict[str, Any]) -> bytes:
    """Generate a PDF CAM report and return the bytes."""
    doc = Document()
    
    # ══════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════
    title = doc.add_heading('Credit Appraisal Memorandum', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('FedCredit Score — AI Financial Intelligence Engine')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
        
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

    next_section_num = 1
    
    # ══════════════════════════════════════════════════
    # 1. COMPANY DETAILS
    # ══════════════════════════════════════════════════
    company = analysis_result.get("company_info", {})
    doc.add_heading(f"{next_section_num}. Company Details", level=2)
    next_section_num += 1
    
    comp_data = [
        ["Field", "Value"],
        ["GSTIN / CIN", company.get("gstin", "N/A")],
        ["Company Name", company.get("company_name", "N/A")],
        ["Location", company.get("location", "N/A")],
    ]
    _create_styled_table(doc, comp_data)

    # ══════════════════════════════════════════════════
    # 2. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════
    llm_summary = analysis_result.get("llm_executive_summary")
    if llm_summary:
        doc.add_heading(f"{next_section_num}. Executive Summary", level=2)
        next_section_num += 1
        
        doc.add_paragraph("AI-generated executive summary based on the provided company data and financial metrics:")
        for para in llm_summary.split("\n\n"):
            if para.strip():
                _add_markdown_para(doc, para.strip())

    # ══════════════════════════════════════════════════
    # LLM CAM NARRATIVE SECTIONS (lazily generated if needed)
    # ══════════════════════════════════════════════════
    cam = analysis_result.get("llm_cam_sections")
    if not cam:
        try:
            from backend.intelligence_layer.llm_service import generate_cam_narrative
            cam = generate_cam_narrative(analysis_result)
        except Exception:
            cam = None

    if cam and isinstance(cam, dict):
        cam_sections = [
            ("Borrower Profile", "borrower_profile"),
            ("Industry & Market Analysis", "industry_analysis"),
            ("Financial Analysis — Narrative", "financial_analysis"),
            ("Risk Assessment — Narrative", "risk_assessment"),
            ("Credit Evaluation Framework", "credit_evaluation"),
            ("Loan Recommendation — Assessment", "loan_recommendation"),
            ("Final Credit Perspective", "final_credit_perspective"),
        ]
        for section_title, key in cam_sections:
            text = cam.get(key)
            if text and isinstance(text, str) and text.strip():
                doc.add_heading(f"{next_section_num}. {section_title}", level=2)
                next_section_num += 1
                for para in text.split("\n\n"):
                    if para.strip():
                        _add_markdown_para(doc, para.strip())

    # ══════════════════════════════════════════════════
    # OFFICER ASSESSMENT
    # ══════════════════════════════════════════════════
    doc.add_heading(f"{next_section_num}. Officer Assessment", level=2)
    next_section_num += 1
    officer_insights = company.get("insights", "") or ""
    
    doc.add_paragraph("The following assessment was provided by the credit officer after reviewing the company's documents, operations, and management:")
    if officer_insights.strip():
        p = doc.add_paragraph(f'"{officer_insights.strip()}"')
        for run in p.runs:
            run.italic = True
    else:
        doc.add_paragraph("No officer insights were provided.")

    # ══════════════════════════════════════════════════
    # FEDERATED CREDIT SCORE
    # ══════════════════════════════════════════════════
    scoring = analysis_result.get("scoring", {})
    final_score = scoring.get("federated_score", 0)
    original_score = scoring.get("original_score", final_score)
    insight_adj = scoring.get("insight_adjustment", 0)
    risk = scoring.get("risk_category", "N/A")
    
    doc.add_heading(f"{next_section_num}. Federated Credit Score", level=2)
    next_section_num += 1
    
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run1 = score_p.add_run(f"{final_score} / 1000\n")
    s_run1.bold = True
    s_run1.font.size = Pt(24)
    s_run2 = score_p.add_run(risk)
    s_run2.bold = True
    s_run2.font.size = Pt(14)

    doc.add_heading("Score Impact of Officer Assessment", level=3)
    score_data = [
        ["Metric", "Value"],
        ["Score Before Officer Review", str(original_score)],
        ["Officer Insight Adjustment", f"{'+' if insight_adj >= 0 else ''}{insight_adj} points"],
        ["Final Federated Score", str(final_score)],
        ["Risk Category", risk]
    ]
    _create_styled_table(doc, score_data)
    
    if insight_adj != 0:
        dir_word = "increased" if insight_adj > 0 else "decreased"
        doc.add_paragraph(f"The officer's assessment {dir_word} the federated score by {abs(insight_adj)} points.")
        
    doc.add_paragraph(
        f"A score of {final_score}/1000 indicates {'strong creditworthiness' if final_score >= 700 else 'moderate risk' if final_score >= 400 else 'significant risk'}."
    )

    # ══════════════════════════════════════════════════
    # 5Cs OF CREDIT ANALYSIS
    # ══════════════════════════════════════════════════
    five_cs = scoring.get("five_cs", [])
    if five_cs:
        doc.add_heading(f"{next_section_num}. 5Cs of Credit Analysis", level=2)
        next_section_num += 1
        for c in five_cs:
            doc.add_heading(f"{c.get('name', 'N/A')} — Score: {c.get('score', 0)}/100", level=3)
            doc.add_paragraph(f"{c.get('description', '')}")
            _add_markdown_para(doc, c.get("explanation", ""))
            
            factors = c.get("factors", [])
            for f in factors:
                doc.add_paragraph(f"• {f}")

    # ══════════════════════════════════════════════════
    # RISK SUMMARY NARRATIVE
    # ══════════════════════════════════════════════════
    risk_narrative = scoring.get("risk_narrative", "")
    if risk_narrative:
        doc.add_heading(f"{next_section_num}. Risk Assessment Summary", level=2)
        next_section_num += 1
        
        for para in risk_narrative.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            if ":" in para and para.split(":")[0].isupper():
                header_part, _, body_part = para.partition(":")
                doc.add_heading(header_part.strip(), level=3)
                if body_part.strip():
                    _add_markdown_para(doc, body_part.strip())
            else:
                _add_markdown_para(doc, para)

    # ══════════════════════════════════════════════════
    # BANK ASSESSMENTS
    # ══════════════════════════════════════════════════
    bank_scores = scoring.get("bank_scores", [])
    if bank_scores:
        doc.add_heading(f"{next_section_num}. Individual Bank Assessments", level=2)
        next_section_num += 1
        
        bank_data = [["Bank", "Focus Area", "Score"]]
        for b in bank_scores:
            bank_data.append([b.get("bank_name", ""), b.get("focus", ""), str(b.get("score", 0))])
        _create_styled_table(doc, bank_data)
        
        bank_summary = scoring.get("bank_summary", "")
        if bank_summary:
            doc.add_heading("Bank Assessment Summary", level=3)
            _add_markdown_para(doc, bank_summary)

    # ══════════════════════════════════════════════════
    # END-OF-REPORT SECTIONS
    # ══════════════════════════════════════════════════
    
    # Financial Metrics
    fin = analysis_result.get("financial_data", {})
    if fin:
        doc.add_heading(f"{next_section_num}. Extracted Financial Metrics", level=2)
        next_section_num += 1
        fin_rows = [["Metric", "Value"]]
        for key in ["turnover", "debt_ratio", "profit_margin", "capacity_utilization", "total_assets", "total_liabilities"]:
            if key in fin:
                label = key.replace("_", " ").title()
                fin_rows.append([label, str(fin[key])])
        if len(fin_rows) > 1:
            _create_styled_table(doc, fin_rows)

    # Intelligence Signals
    intel = analysis_result.get("intelligence", {})
    if intel:
        doc.add_heading(f"{next_section_num}. External Intelligence Signals", level=2)
        next_section_num += 1
        intel_rows = [["Signal", "Value"]]
        intel_rows.append(["News Sentiment", f"{intel.get('news_sentiment', 'N/A')} / 10"])
        intel_rows.append(["MCA Compliance", f"{intel.get('mca_compliance', 'N/A')} / 10"])
        intel_rows.append(["Court Cases", str(intel.get('court_cases', 'N/A'))])
        intel_rows.append(["NDVI Activity", str(intel.get('ndvi_activity', 'N/A'))])
        _create_styled_table(doc, intel_rows)

    # Loan Recommendation
    loan = scoring.get("loan_recommendation", {})
    if loan:
        doc.add_heading(f"{next_section_num}. Loan Recommendation", level=2)
        next_section_num += 1
        loan_rows = [["Parameter", "Value"]]
        loan_rows.append(["Recommended Loan", loan.get("recommended_loan", "N/A")])
        loan_rows.append(["Interest Rate", loan.get("interest_rate", "N/A")])
        loan_rows.append(["Tenure", loan.get("tenure", "N/A")])
        loan_rows.append(["Likelihood", loan.get("approval_likelihood", "N/A")])
        _create_styled_table(doc, loan_rows)
        _add_markdown_para(doc, loan.get("explanation", ""))

    # Enhanced Financial Ratios
    fr = scoring.get("financial_ratios", {})
    if fr:
        doc.add_heading(f"{next_section_num}. Enhanced Financial Ratios", level=2)
        next_section_num += 1
        r_rows = [["Metric", "Value", "Assessment"]]
        for key in ["debt_ebitda", "dscr", "net_profit_margin", "debt_equity"]:
            r = fr.get(key, {})
            if isinstance(r, dict) and "value" in r:
                r_rows.append([key.replace("_", " ").title(), str(r.get("value", "N/A")), str(r.get("assessment", "N/A"))])
        if len(r_rows) > 1:
            _create_styled_table(doc, r_rows)

    # Detailed Risk Assessment
    dr = scoring.get("detailed_risk_assessment", [])
    if dr:
        doc.add_heading(f"{next_section_num}. Detailed Risk Assessment", level=2)
        next_section_num += 1
        dr_rows = [["Category", "Severity", "Description"]]
        for r in dr:
            dr_rows.append([r.get("category", ""), r.get("severity", ""), str(r.get("description", ""))[:150]])
        _create_styled_table(doc, dr_rows)
        for r in dr:
            if r.get("severity") in ["Medium", "High"] and r.get("mitigation"):
                doc.add_paragraph(f"• Mitigation ({r['category']}): {r['mitigation']}")

    # Conclusion
    doc.add_heading(f"{next_section_num}. Conclusion", level=2)
    next_section_num += 1
    doc.add_paragraph(f"Based on the comprehensive credit appraisal, the company receives a federated score of {final_score}/1000 ({risk}).")

    # Footer
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph("Generated by FedCredit Score — AI Financial Intelligence Engine")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ══════════════════════════════════════════════════
    # DOCX TO PDF CONVERSION
    # ══════════════════════════════════════════════════
    try:
        from docx2pdf import convert
        
        # We need absolute file paths
        temp_dir = tempfile.gettempdir()
        base_name = str(uuid.uuid4())
        
        docx_path = os.path.join(temp_dir, f"{base_name}.docx")
        pdf_path = os.path.join(temp_dir, f"{base_name}.pdf")
        
        # Save DOCX
        doc.save(docx_path)
        
        try:
            # Convert DOCX to PDF using COM object (requires Word on Windows)
            convert(docx_path, pdf_path)
            
            # Read back as bytes
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
                
            return pdf_bytes
            
        finally:
            # Cleanup temp files
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                
    except ImportError:
        logger.error("docx2pdf is not installed.")
        raise RuntimeError("docx2pdf package is required for PDF generation.")
    except Exception as e:
        logger.error(f"Failed to generate PDF from DOCX: {e}")
        # If Word isn't installed or conversion fails, we could return DOCX bytes as a fallback
        # by saving into BytesIO, but to match the previous requirement we raise an error.
        raise RuntimeError(f"PDF generation failed: {e}")
