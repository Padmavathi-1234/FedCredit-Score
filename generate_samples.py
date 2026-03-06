"""
Generate sample financial documents for demo/testing purposes.
Run this script to create sample_documents/*.
"""

import os
import sys

# Add project root to path
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DIR = os.path.join(PROJECT_ROOT, "sample_documents")
os.makedirs(SAMPLE_DIR, exist_ok=True)


def create_sample_xlsx():
    """Create a sample balance sheet Excel file."""
    try:
        import openpyxl
    except ImportError:
        print("openpyxl not installed, skipping XLSX generation")
        return

    wb = openpyxl.Workbook()

    # Sheet 1: Balance Sheet
    ws = wb.active
    ws.title = "Balance Sheet"

    headers = ["Particulars", "FY 2024-25 (₹)", "FY 2023-24 (₹)"]
    ws.append(headers)
    ws.append([])

    ws.append(["ASSETS", "", ""])
    ws.append(["Fixed Assets", 15000000, 13500000])
    ws.append(["Current Assets", 22000000, 19800000])
    ws.append(["Investments", 8000000, 7200000])
    ws.append(["Cash & Bank Balance", 5000000, 4500000])
    ws.append(["Total Assets", 50000000, 45000000])
    ws.append([])

    ws.append(["LIABILITIES", "", ""])
    ws.append(["Share Capital", 10000000, 10000000])
    ws.append(["Reserves & Surplus", 12000000, 10000000])
    ws.append(["Long Term Borrowings", 8000000, 7500000])
    ws.append(["Current Liabilities", 15000000, 12500000])
    ws.append(["Short Term Borrowings", 5000000, 5000000])
    ws.append(["Total Liabilities", 50000000, 45000000])
    ws.append([])

    ws.append(["KEY RATIOS", "", ""])
    ws.append(["Total Revenue / Turnover", 52000000, 47000000])
    ws.append(["Net Profit Margin", "18%", "16%"])
    ws.append(["Debt to Equity Ratio", 0.42, 0.45])
    ws.append(["Capacity Utilization", "80%", "75%"])
    ws.append(["Return on Equity", "22%", "19%"])

    # Sheet 2: P&L Statement
    ws2 = wb.create_sheet("Profit & Loss")
    ws2.append(["Particulars", "FY 2024-25 (₹)", "FY 2023-24 (₹)"])
    ws2.append([])
    ws2.append(["Gross Revenue", 55000000, 50000000])
    ws2.append(["Less: Excise Duty", 3000000, 3000000])
    ws2.append(["Net Revenue / Turnover", 52000000, 47000000])
    ws2.append(["Cost of Goods Sold", 32000000, 29500000])
    ws2.append(["Gross Profit", 20000000, 17500000])
    ws2.append(["Operating Expenses", 8000000, 7200000])
    ws2.append(["EBITDA", 12000000, 10300000])
    ws2.append(["Depreciation", 1500000, 1400000])
    ws2.append(["Interest", 1140000, 1050000])
    ws2.append(["Profit Before Tax", 9360000, 7850000])
    ws2.append(["Tax", 2340000, 1962000])
    ws2.append(["Net Profit", 7020000, 5888000])

    # Auto-width
    for sheet in [ws, ws2]:
        for col in sheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in col)
            sheet.column_dimensions[col[0].column_letter].width = max_length + 4

    path = os.path.join(SAMPLE_DIR, "sample_balance_sheet.xlsx")
    wb.save(path)
    print(f"Created: {path}")


def create_sample_docx():
    """Create a sample financial report Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed, skipping DOCX generation")
        return

    doc = Document()

    # Title
    title = doc.add_heading("Annual Financial Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Acme Industries Pvt. Ltd.")
    doc.add_paragraph("Financial Year 2024-25")
    doc.add_paragraph("GSTIN: 29ABCDE1234F1Z5")
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Acme Industries Pvt. Ltd. has demonstrated strong financial performance "
        "in FY 2024-25. The company reported a total turnover of ₹52,00,00,000 "
        "(₹52 Crore), reflecting a year-over-year growth of 10.6%. The net profit "
        "margin stands at 18%, indicating healthy operational efficiency."
    )
    doc.add_paragraph(
        "The debt-to-equity ratio of 0.42 remains within acceptable limits, "
        "suggesting prudent financial management. Capacity utilization has improved "
        "to 80% from 75% in the previous year, driven by new production line additions."
    )

    # Financial Highlights
    doc.add_heading("Financial Highlights", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "FY 2024-25"
    hdr[2].text = "FY 2023-24"

    data = [
        ("Total Revenue", "₹52,00,00,000", "₹47,00,00,000"),
        ("Net Profit", "₹7,02,00,000", "₹5,88,80,000"),
        ("Net Profit Margin", "18%", "16%"),
        ("Debt Ratio", "0.42", "0.45"),
        ("Capacity Utilization", "80%", "75%"),
        ("Total Assets", "₹50,00,00,000", "₹45,00,00,000"),
        ("Return on Equity", "22%", "19%"),
    ]

    for metric, fy25, fy24 in data:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = fy25
        row[2].text = fy24

    doc.add_paragraph("")

    # Audit Notes
    doc.add_heading("Audit Notes", level=1)
    doc.add_paragraph(
        "The audit of the financial statements for FY 2024-25 was conducted by "
        "M/s Sharma & Associates, Chartered Accountants. The audit report is "
        "unqualified with no material observations. All disclosures are in "
        "compliance with the Companies Act, 2013 and applicable Indian "
        "Accounting Standards (Ind AS)."
    )
    doc.add_paragraph(
        "The company has maintained proper books of accounts and the financial "
        "statements give a true and fair view of the state of affairs. "
        "No instances of fraud have been reported."
    )

    # Operational Overview
    doc.add_heading("Operational Overview", level=1)
    doc.add_paragraph(
        "The company operates three manufacturing units across Karnataka and "
        "Tamil Nadu. The combined production capacity utilization reached 80% "
        "during the review period. An expansion project to add a fourth unit "
        "is expected to commence in Q1 FY 2025-26."
    )

    # Risk Factors
    doc.add_heading("Risk Factors", level=1)
    doc.add_paragraph(
        "Key risks include raw material price volatility, regulatory changes in "
        "the manufacturing sector, and potential disruptions from economic "
        "slowdown. The company has implemented hedging strategies and maintains "
        "adequate insurance coverage to mitigate these risks."
    )

    path = os.path.join(SAMPLE_DIR, "sample_financial_report.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_sample_pdf_text():
    """Create a simple text-based PDF (using reportlab if available)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
    except ImportError:
        print("reportlab not installed, skipping PDF generation")
        return

    path = os.path.join(SAMPLE_DIR, "sample_audit_report.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()

    elements = []
    elements.append(Paragraph("Audit Report — Acme Industries Pvt. Ltd.", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("Financial Year 2024-25", styles["Heading2"]))
    elements.append(Spacer(1, 12))

    paragraphs = [
        "We have audited the accompanying financial statements of Acme Industries "
        "Pvt. Ltd., which comprise the Balance Sheet as at March 31, 2025, and the "
        "Statement of Profit and Loss for the year then ended.",

        "Total Revenue / Turnover: ₹52,00,00,000 (Rs. 52 Crore).",

        "The company reported a net profit margin of 18% for the financial year. "
        "The debt to equity ratio stands at 0.42, reflecting conservative leverage.",

        "Capacity utilization during the year was at 80%, up from 75% in the "
        "previous financial year.",

        "In our opinion, the financial statements give a true and fair view of "
        "the financial position of the Company. The audit report is unqualified. "
        "All statutory filings have been completed on time.",

        "Key Financial Ratios:",
        "Turnover: 52000000",
        "Debt Ratio: 0.42",
        "Profit Margin: 18%",
        "Capacity Utilization: 80%",
        "Return on Equity: 22%",
    ]

    for p in paragraphs:
        elements.append(Paragraph(p, styles["Normal"]))
        elements.append(Spacer(1, 8))

    doc.build(elements)
    print(f"Created: {path}")


if __name__ == "__main__":
    print("Generating sample documents...")
    create_sample_xlsx()
    create_sample_docx()
    create_sample_pdf_text()
    print("Done!")
